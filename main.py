from flask import Flask, render_template, request, redirect, url_for, send_file
from docx import Document
import xml.etree.ElementTree as ET

app = Flask(__name__, template_folder='/home/marilia/Projetos/xml_chico')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files['xml_file']
    print(file.filename)
    tree = ET.parse(file)
    root = tree.getroot()
    document = Document()
    for element in root.iter():
        print(element.text)
        document.add_paragraph(element.text)
    document.save('converted.docx')
    return redirect(url_for('download'))
    # xml_file = request.files['xml_file']
    # # Converta o arquivo XML em um objeto Document do python-docx
    # doc = Document(xml_file)
    # # Salve o objeto Document em um arquivo DOCX
    # doc.save('converted.docx')
    # # Redirecione o usuário para a página de download do arquivo convertido
    # return redirect(url_for('download'))

@app.route('/download')
def download():
    # Envie o arquivo convertido para o usuário
    
    return send_file('converted.docx', as_attachment=True)

if __name__ == '__main__':
    app.run()

